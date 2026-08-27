import os
import io
import socket
import sqlite3
import urllib.parse
import qrcode
import docx
from fastapi import FastAPI, Request, Form, Response, Cookie
from fastapi.responses import HTMLResponse, RedirectResponse, FileResponse
from jinja2 import Environment, FileSystemLoader
from database import get_db, init_db

app = FastAPI(title="Sistema ExpoTecnica MEP")

templates_env = Environment(loader=FileSystemLoader("templates"))

# Clave de Administrador
ADMIN_PASSWORD = "admin123"


@app.on_event("startup")
def startup_event():
    init_db()


def obtener_institucion():
    conn = get_db()
    conf = conn.execute("SELECT nombre_institucion FROM configuracion WHERE id = 1").fetchone()
    conn.close()
    return conf["nombre_institucion"] if conf else "Colegio Técnico Profesional"


def obtener_url_base(request: Request) -> str:
    host = request.headers.get("x-forwarded-host") or request.headers.get("host") or ""
    if "onrender.com" in host or (not host.startswith("127.") and not host.startswith("192.") and not host.startswith(
            "10.") and not host.startswith("localhost")):
        scheme = "https"
    else:
        scheme = request.headers.get("x-forwarded-proto") or request.url.scheme or "http"
    if host:
        return f"{scheme}://{host}"
    return "https://sistema-expotecnica.onrender.com"


def es_admin_autenticado(admin_session: str = None) -> bool:
    return admin_session == "sesion_activa_admin"


def reemplazar_etiquetas_word(doc, reemplazos):
    for p in doc.paragraphs:
        for tag, valor in reemplazos.items():
            if tag in p.text:
                p.text = p.text.replace(tag, str(valor))

    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for p in celda.paragraphs:
                    for tag, valor in reemplazos.items():
                        if tag in p.text:
                            p.text = p.text.replace(tag, str(valor))


# ----------------- LOGIN Y SEGURIDAD ADMIN -----------------
@app.get("/login", response_class=HTMLResponse)
def vista_login(error: str = None):
    template = templates_env.get_template("login.html")
    return template.render(error=error)


@app.post("/login")
def procesar_login(password: str = Form(...)):
    if password == ADMIN_PASSWORD:
        response = RedirectResponse(url="/", status_code=303)
        response.set_cookie(key="admin_session", value="sesion_activa_admin", httponly=True)
        return response
    else:
        err = urllib.parse.quote("Contraseña de administrador incorrecta.")
        return RedirectResponse(url=f"/login?error={err}", status_code=303)


@app.get("/logout")
def procesar_logout():
    response = RedirectResponse(url="/login", status_code=303)
    response.delete_cookie(key="admin_session")
    return response


# ----------------- PANEL GENERAL (ADMIN) -----------------
@app.get("/", response_class=HTMLResponse)
def index(request: Request, msg: str = None, error: str = None, admin_session: str = Cookie(default=None)):
    if not es_admin_autenticado(admin_session):
        return RedirectResponse(url="/login", status_code=303)

    conn = get_db()
    jueces = conn.execute("SELECT * FROM jueces ORDER BY nombre ASC").fetchall()
    proyectos = conn.execute("SELECT * FROM proyectos ORDER BY codigo ASC").fetchall()

    proyectos_summary = []
    for p in proyectos:
        total_evaluados = conn.execute(
            "SELECT COUNT(*) FROM asignaciones WHERE proyecto_id = ? AND estado = 'EVALUADO'",
            (p['id'],)
        ).fetchone()[0]

        promedio = conn.execute('''
            SELECT AVG(e.nota_final) 
            FROM evaluaciones e 
            JOIN asignaciones a ON e.asignacion_id = a.id 
            WHERE a.proyecto_id = ?
        ''', (p['id'],)).fetchone()[0]

        proyectos_summary.append({
            "id": p["id"],
            "codigo": p["codigo"],
            "nombre": p["nombre"],
            "especialidad": p["especialidad"],
            "estudiantes": [p["estudiante_1"], p["estudiante_2"], p["estudiante_3"]],
            "total_evaluados": total_evaluados,
            "promedio": round(promedio, 2) if promedio is not None else "Sin notas"
        })
    conn.close()

    nombre_institucion = obtener_institucion()
    base_url = obtener_url_base(request)
    template = templates_env.get_template("index.html")
    return template.render(
        proyectos=proyectos_summary,
        jueces=jueces,
        base_url=base_url,
        nombre_institucion=nombre_institucion,
        msg=msg,
        error=error
    )


# ----------------- 1. GANADORES Y RANKING (EXCLUSIVO ADMIN) -----------------
@app.get("/admin/ganadores", response_class=HTMLResponse)
def vista_ganadores(admin_session: str = Cookie(default=None)):
    if not es_admin_autenticado(admin_session):
        return RedirectResponse(url="/login", status_code=303)

    conn = get_db()
    proyectos = conn.execute("SELECT * FROM proyectos").fetchall()

    ranking = []
    for p in proyectos:
        asigs = conn.execute('''
            SELECT a.estado, e.nota_final 
            FROM asignaciones a
            LEFT JOIN evaluaciones e ON e.asignacion_id = a.id
            WHERE a.proyecto_id = ?
        ''', (p['id'],)).fetchall()

        evaluadas = [a['nota_final'] for a in asigs if a['nota_final'] is not None]
        total_evaluadas = len(evaluadas)
        promedio = round(sum(evaluadas) / total_evaluadas, 2) if total_evaluadas > 0 else 0.0

        estudiantes = [p['estudiante_1']]
        if p['estudiante_2']: estudiantes.append(p['estudiante_2'])
        if p['estudiante_3']: estudiantes.append(p['estudiante_3'])

        ranking.append({
            "id": p["id"],
            "codigo": p["codigo"],
            "nombre": p["nombre"],
            "estudiantes": estudiantes,
            "total_evaluadas": total_evaluadas,
            "promedio": promedio,
            "completado": (total_evaluadas == 3)
        })
    conn.close()

    # Ordenar por mayor promedio descendente
    ranking.sort(key=lambda x: x["promedio"], reverse=True)

    template = templates_env.get_template("ganadores.html")
    return template.render(ranking=ranking, nombre_institucion=obtener_institucion())


# ----------------- 2. CONFIGURACIÓN DE INSTITUCIÓN -----------------
@app.post("/admin/configuracion/guardar")
def guardar_configuracion(nombre_institucion: str = Form(...), admin_session: str = Cookie(default=None)):
    if not es_admin_autenticado(admin_session):
        return RedirectResponse(url="/login", status_code=303)

    conn = get_db()
    conn.execute("UPDATE configuracion SET nombre_institucion = ? WHERE id = 1", (nombre_institucion.strip(),))
    conn.commit()
    conn.close()

    msg = urllib.parse.quote("Nombre de la institución actualizado correctamente.")
    return RedirectResponse(url=f"/?msg={msg}", status_code=303)


# ----------------- 4. REINICIAR BASE DE DATOS A CERO -----------------
@app.post("/admin/reset-total")
def reset_base_datos(confirmacion: str = Form(...), admin_session: str = Cookie(default=None)):
    if not es_admin_autenticado(admin_session):
        return RedirectResponse(url="/login", status_code=303)

    if confirmacion.strip().upper() != "BORRAR TODO":
        err = urllib.parse.quote("Debe escribir exactamente 'BORRAR TODO' para confirmar.")
        return RedirectResponse(url=f"/?error={err}", status_code=303)

    conn = get_db()
    cursor = conn.cursor()
    cursor.execute("DELETE FROM evaluaciones")
    cursor.execute("DELETE FROM asignaciones")
    cursor.execute("DELETE FROM proyectos")
    cursor.execute("DELETE FROM jueces")
    conn.commit()
    conn.close()

    msg = urllib.parse.quote("Se ha restablecido la base de datos a cero exitosamente.")
    return RedirectResponse(url=f"/?msg={msg}", status_code=303)


# ----------------- GENERADOR QR -----------------
@app.get("/qr/juez/{juez_id}.png")
def generar_qr_juez(request: Request, juez_id: int):
    base_url = obtener_url_base(request)
    url_juez = f"{base_url}/juez/panel?juez_id={juez_id}"

    qr = qrcode.QRCode(version=1, box_size=8, border=2)
    qr.add_data(url_juez)
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white")

    img_byte_arr = io.BytesIO()
    img.save(img_byte_arr, format='PNG')
    return Response(content=img_byte_arr.getvalue(), media_type="image/png")


# ----------------- PROYECTOS (ADMIN) -----------------
@app.get("/admin/proyectos", response_class=HTMLResponse)
def vista_proyectos(error: str = None, admin_session: str = Cookie(default=None)):
    if not es_admin_autenticado(admin_session):
        return RedirectResponse(url="/login", status_code=303)

    conn = get_db()
    proyectos = conn.execute("SELECT * FROM proyectos ORDER BY id DESC").fetchall()
    jueces = conn.execute("SELECT * FROM jueces ORDER BY nombre ASC").fetchall()

    proyectos_data = []
    for p in proyectos:
        asigs = conn.execute('''
            SELECT a.id as asignacion_id, a.estado, j.id as juez_id, j.nombre as juez_nombre 
            FROM asignaciones a 
            JOIN jueces j ON a.juez_id = j.id 
            WHERE a.proyecto_id = ?
        ''', (p['id'],)).fetchall()
        proyectos_data.append({"proyecto": p, "asignaciones": asigs})
    conn.close()

    template = templates_env.get_template("admin_proyectos.html")
    return template.render(proyectos_data=proyectos_data, jueces=jueces, error=error)


@app.post("/admin/proyectos/crear")
def crear_proyecto(
        codigo: str = Form(...),
        nombre: str = Form(...),
        especialidad: str = Form("STEAM"),
        estudiante_1: str = Form(...),
        estudiante_2: str = Form(""),
        estudiante_3: str = Form(""),
        juez_1: int = Form(None),
        juez_2: int = Form(None),
        juez_3: int = Form(None),
        admin_session: str = Cookie(default=None)
):
    if not es_admin_autenticado(admin_session):
        return RedirectResponse(url="/login", status_code=303)

    if not juez_1 or not juez_2 or not juez_3:
        err = urllib.parse.quote("Debe seleccionar los 3 jueces obligatorios.")
        return RedirectResponse(url=f"/admin/proyectos?error={err}", status_code=303)

    jueces_seleccionados = {juez_1, juez_2, juez_3}
    if len(jueces_seleccionados) < 3:
        err = urllib.parse.quote("Los 3 jueces asignados deben ser personas diferentes.")
        return RedirectResponse(url=f"/admin/proyectos?error={err}", status_code=303)

    conn = get_db()
    cursor = conn.cursor()
    try:
        cursor.execute('''
            INSERT INTO proyectos (codigo, nombre, especialidad, seccion, estudiante_1, estudiante_2, estudiante_3, descripcion)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        ''', (codigo, nombre, especialidad, "", estudiante_1, estudiante_2, estudiante_3, ""))
        proyecto_id = cursor.lastrowid

        for j_id in jueces_seleccionados:
            cursor.execute("INSERT INTO asignaciones (proyecto_id, juez_id) VALUES (?, ?)", (proyecto_id, j_id))

        conn.commit()
    except sqlite3.IntegrityError:
        conn.rollback()
        err = urllib.parse.quote(f"El código de proyecto '{codigo}' ya existe.")
        return RedirectResponse(url=f"/admin/proyectos?error={err}", status_code=303)
    finally:
        conn.close()

    return RedirectResponse(url="/admin/proyectos", status_code=303)


@app.post("/admin/proyectos/eliminar")
def eliminar_proyecto(proyecto_id: int = Form(...), admin_session: str = Cookie(default=None)):
    if not es_admin_autenticado(admin_session):
        return RedirectResponse(url="/login", status_code=303)

    conn = get_db()
    cursor = conn.cursor()
    cursor.execute("DELETE FROM proyectos WHERE id = ?", (proyecto_id,))
    conn.commit()
    conn.close()
    return RedirectResponse(url="/admin/proyectos", status_code=303)


@app.post("/admin/asignacion/cambiar-juez")
def cambiar_juez_asignacion(
        asignacion_id: int = Form(...),
        nuevo_juez_id: int = Form(...),
        admin_session: str = Cookie(default=None)
):
    if not es_admin_autenticado(admin_session):
        return RedirectResponse(url="/login", status_code=303)

    conn = get_db()
    cursor = conn.cursor()

    asig = conn.execute("SELECT estado, proyecto_id FROM asignaciones WHERE id = ?", (asignacion_id,)).fetchone()
    if asig and asig["estado"] == "EVALUADO":
        conn.close()
        err = urllib.parse.quote("No se puede cambiar el juez porque este proyecto ya fue evaluado.")
        return RedirectResponse(url=f"/admin/proyectos?error={err}", status_code=303)

    existente = conn.execute(
        "SELECT id FROM asignaciones WHERE proyecto_id = ? AND juez_id = ?",
        (asig["proyecto_id"], nuevo_juez_id)
    ).fetchone()

    if existente:
        conn.close()
        err = urllib.parse.quote("El juez seleccionado ya forma parte de la terna de este proyecto.")
        return RedirectResponse(url=f"/admin/proyectos?error={err}", status_code=303)

    cursor.execute("UPDATE asignaciones SET juez_id = ? WHERE id = ?", (nuevo_juez_id, asignacion_id))
    conn.commit()
    conn.close()
    return RedirectResponse(url="/admin/proyectos", status_code=303)


# ----------------- JUECES (ADMIN) -----------------
@app.get("/admin/jueces", response_class=HTMLResponse)
def vista_jueces(error: str = None, admin_session: str = Cookie(default=None)):
    if not es_admin_autenticado(admin_session):
        return RedirectResponse(url="/login", status_code=303)

    conn = get_db()
    jueces = conn.execute("SELECT * FROM jueces ORDER BY nombre ASC").fetchall()
    conn.close()
    template = templates_env.get_template("admin_jueces.html")
    return template.render(jueces=jueces, error=error)


@app.post("/admin/jueces/crear")
def crear_juez(
        cedula: str = Form(...),
        nombre: str = Form(...),
        email: str = Form(""),
        especialidad: str = Form("STEAM"),
        telefono: str = Form("")
):
    conn = get_db()
    cursor = conn.cursor()
    try:
        cursor.execute('''
            INSERT INTO jueces (cedula, nombre, email, especialidad, telefono)
            VALUES (?, ?, ?, ?, ?)
        ''', (cedula, nombre, email, especialidad, telefono))
        conn.commit()
    except sqlite3.IntegrityError:
        conn.rollback()
        err = urllib.parse.quote(f"Ya existe un juez registrado con la cédula '{cedula}'.")
        return RedirectResponse(url=f"/admin/jueces?error={err}", status_code=303)
    finally:
        conn.close()

    return RedirectResponse(url="/admin/jueces", status_code=303)


@app.post("/admin/jueces/eliminar")
def eliminar_juez(juez_id: int = Form(...), admin_session: str = Cookie(default=None)):
    if not es_admin_autenticado(admin_session):
        return RedirectResponse(url="/login", status_code=303)

    conn = get_db()
    cursor = conn.cursor()
    cursor.execute("DELETE FROM jueces WHERE id = ?", (juez_id,))
    conn.commit()
    conn.close()
    return RedirectResponse(url="/admin/jueces", status_code=303)


# ----------------- PORTAL Y EVALUACIÓN DEL JUEZ -----------------
@app.get("/juez/panel", response_class=HTMLResponse)
def juez_panel(juez_id: int):
    conn = get_db()
    juez = conn.execute("SELECT * FROM jueces WHERE id = ?", (juez_id,)).fetchone()

    if not juez:
        conn.close()
        return HTMLResponse("<h2>Error: Evaluador no encontrado.</h2>", status_code=404)

    asignaciones = conn.execute('''
        SELECT a.id as asignacion_id, a.estado, p.codigo, p.nombre, p.especialidad, 
               p.estudiante_1, p.estudiante_2, p.estudiante_3,
               e.nota_final, e.id as evaluacion_id
        FROM asignaciones a
        JOIN proyectos p ON a.proyecto_id = p.id
        LEFT JOIN evaluaciones e ON e.asignacion_id = a.id
        WHERE a.juez_id = ?
        ORDER BY a.estado ASC, p.codigo ASC
    ''', (juez_id,)).fetchall()
    conn.close()

    template = templates_env.get_template("juez_panel.html")
    return template.render(juez=juez, asignaciones=asignaciones)


@app.get("/evaluar/{asignacion_id}", response_class=HTMLResponse)
def formulario_evaluar(asignacion_id: int):
    conn = get_db()
    asig = conn.execute('''
        SELECT a.id as asignacion_id, a.estado, 
               j.id as juez_id, j.nombre as juez_nombre, j.cedula as juez_cedula,
               p.id as proyecto_id, p.codigo, p.nombre as proyecto_nombre, 
               p.especialidad, p.estudiante_1, p.estudiante_2, p.estudiante_3
        FROM asignaciones a
        JOIN jueces j ON a.juez_id = j.id
        JOIN proyectos p ON a.proyecto_id = p.id
        WHERE a.id = ?
    ''', (asignacion_id,)).fetchone()
    conn.close()

    if not asig:
        return HTMLResponse("<h2>Asignación no encontrada.</h2>", status_code=404)

    template = templates_env.get_template("formulario_evaluacion.html")
    return template.render(asig=asig)


@app.post("/evaluar/{asignacion_id}/guardar")
async def guardar_evaluacion(request: Request, asignacion_id: int):
    form_data = await request.form()

    campos = [
        'i_a', 'i_b', 'i_c', 'i_d', 'i_e',
        'ii_a', 'ii_b', 'ii_c', 'ii_d', 'ii_e', 'ii_f',
        'iii_a', 'iii_b', 'iii_c', 'iii_d', 'iii_e', 'iii_f', 'iii_g', 'iii_h',
        'iv_a', 'iv_b', 'iv_c', 'iv_d', 'iv_e', 'iv_f', 'iv_g', 'iv_h', 'iv_i', 'iv_j',
        'v_a_inf', 'v_b_inf', 'v_c_inf', 'v_a_bit', 'v_b_bit', 'v_c_bit', 'v_a_car', 'v_b_car'
    ]

    valores = []
    puntaje_total = 0
    for campo in campos:
        val = int(form_data.get(campo, 0))
        valores.append(val)
        puntaje_total += val

    observaciones = form_data.get("observaciones", "")
    recomendaciones = form_data.get("recomendaciones", "")

    # 3) CÁLCULO MATEMÁTICO EXACTO: 37 indicadores * 3 pts = 111 pts máx = 100%
    nota_final = round((puntaje_total * 100.0) / 111.0, 2)

    conn = get_db()
    cursor = conn.cursor()
    try:
        placeholders = ', '.join(['?'] * len(campos))
        columnas = ', '.join(campos)
        query = f'''
            INSERT INTO evaluaciones (
                asignacion_id, {columnas}, puntaje_total, nota_final, observaciones, recomendaciones
            ) VALUES (?, {placeholders}, ?, ?, ?, ?)
        '''
        datos_insert = [asignacion_id] + valores + [puntaje_total, nota_final, observaciones, recomendaciones]
        cursor.execute(query, datos_insert)
        cursor.execute("UPDATE asignaciones SET estado = 'EVALUADO' WHERE id = ?", (asignacion_id,))
        conn.commit()
        eval_id = cursor.lastrowid
    except Exception as e:
        conn.rollback()
        return HTMLResponse(f"<h3>Error al guardar evaluación: {str(e)}</h3>", status_code=400)
    finally:
        juez_row = conn.execute("SELECT juez_id FROM asignaciones WHERE id = ?", (asignacion_id,)).fetchone()
        juez_id = juez_row[0] if juez_row else 1
        conn.close()

    return RedirectResponse(url=f"/evaluacion/exito/{eval_id}?juez_id={juez_id}", status_code=303)


@app.get("/evaluacion/exito/{eval_id}", response_class=HTMLResponse)
def evaluacion_exitosa(eval_id: int, juez_id: int = 1):
    template = templates_env.get_template("evaluacion_exitosa.html")
    return template.render(eval_id=eval_id, juez_id=juez_id)


# ----------------- DESCARGAS (OFICIALES) -----------------
@app.get("/evaluacion/{eval_id}/pdf")
@app.get("/evaluacion/{eval_id}/docx")
def generar_documento_oficial_word(eval_id: int):
    conn = get_db()
    data = conn.execute('''
        SELECT e.*, 
               p.codigo as proyecto_codigo, p.nombre as proyecto_nombre, p.especialidad,
               p.estudiante_1, p.estudiante_2, p.estudiante_3,
               j.nombre as juez_nombre, j.cedula as juez_cedula
        FROM evaluaciones e
        JOIN asignaciones a ON e.asignacion_id = a.id
        JOIN proyectos p ON a.proyecto_id = p.id
        JOIN jueces j ON a.juez_id = j.id
        WHERE e.id = ?
    ''', (eval_id,)).fetchone()
    conn.close()

    if not data:
        return HTMLResponse("Evaluación no encontrada", status_code=404)

    plantilla_path = os.path.join("plantillas", "Plantilla Steam.docx")
    if not os.path.exists(plantilla_path):
        return HTMLResponse("No se encontró 'Plantilla Steam.docx' dentro de la carpeta 'plantillas/'", status_code=500)

    doc = docx.Document(plantilla_path)

    suma_i = data["i_a"] + data["i_b"] + data["i_c"] + data["i_d"] + data["i_e"]
    suma_ii = data["ii_a"] + data["ii_b"] + data["ii_c"] + data["ii_d"] + data["ii_e"] + data["ii_f"]
    suma_iii = (data["iii_a"] + data["iii_b"] + data["iii_c"] + data["iii_d"] +
                data["iii_e"] + data["iii_f"] + data["iii_g"] + data["iii_h"])
    suma_iv = (data["iv_a"] + data["iv_b"] + data["iv_c"] + data["iv_d"] + data["iv_e"] +
               data["iv_f"] + data["iv_g"] + data["iv_h"] + data["iv_i"] + data["iv_j"])
    suma_v = (data["v_a_inf"] + data["v_b_inf"] + data["v_c_inf"] +
              data["v_a_bit"] + data["v_b_bit"] + data["v_c_bit"] +
              data["v_a_car"] + data["v_b_car"])

    reemplazos = {
        "<<Colegio>>": obtener_institucion(),
        "<<Estudiante 1>>": data["estudiante_1"] or "",
        "<<Estudiante 2>>": data["estudiante_2"] or "",
        "<<Estudiante 3>>": data["estudiante_3"] or "",
        "<<Proyecto>>": data["proyecto_nombre"],
        "<<Eje tematico>>": "STEAM",

        "<<Indicador I.a>>": data["i_a"],
        "<<Indicador I.b>>": data["i_b"],
        "<<Indicador I.c>>": data["i_c"],
        "<<Indicador I.d>>": data["i_d"],
        "<<Indicador I.e>>": data["i_e"],
        "<<Suma I>>": suma_i,

        "<<Indicador II.a>>": data["ii_a"],
        "<<Indicador II.b >>": data["ii_b"],
        "<<Indicador II.b>>": data["ii_b"],
        "<<Indicador II.c>>": data["ii_c"],
        "<<Indicador II.d>>": data["ii_d"],
        "<<Indicador II.e>>": data["ii_e"],
        "<<Indicador II.f>>": data["ii_f"],
        "<<Suma II>>": suma_ii,

        "<<Indicador III.a>>": data["iii_a"],
        "<<Indicador III.b>>": data["iii_b"],
        "<<Indicador III.c>>": data["iii_c"],
        "<<Indicador III.d>>": data["iii_d"],
        "<<Indicador III.e>>": data["iii_e"],
        "<<Indicador III.f>>": data["iii_f"],
        "<<Indicador III.g>>": data["iii_g"],
        "<<Indicador III.h>>": data["iii_h"],
        "<<Suma III>>": suma_iii,

        "<<Indicador IV.a>>": data["iv_a"],
        "<<Indicador IV.b>>": data["iv_b"],
        "<<Indicador IV.c>>": data["iv_c"],
        "<<Indicador IV.d>>": data["iv_d"],
        "<<Indicador IV.e>>": data["iv_e"],
        "<<Indicador IV.f>>": data["iv_f"],
        "<<Indicador IV.g>>": data["iv_g"],
        "<<Indicador IV.h>>": data["iv_h"],
        "<<Indicador IV.i >>": data["iv_i"],
        "<<Indicador IV.i>>": data["iv_i"],
        "<<Indicador IV.j >>": data["iv_j"],
        "<<Indicador IV.j>>": data["iv_j"],
        "<<Suma IV>>": suma_iv,

        "<<Indicador V.a (bitacora)>>": data["v_a_bit"],
        "<<Indicador V.b (bitacora))>>": data["v_b_bit"],
        "<<Indicador V.b (bitacora)>>": data["v_b_bit"],
        "<<Indicador V.c (bitacora)>>": data["v_c_bit"],
        "<<Indicador V.a (CP)>>": data["v_a_car"],
        "<<Indicador V.b (CP)>>": data["v_b_car"],
        "<<Suma V>>": suma_v,

        # 3) Sumatoria y Porcentaje sobre 111 puntos
        "<<<Sumatoria>>": f"{data['puntaje_total']} / 111",
        "<<Sumatoria>>": f"{data['puntaje_total']} / 111",
        "<<Calificacion>>": data["nota_final"],
        "<<Porcentaje>>": f"{data['nota_final']}%",
        "<<Recomendaciones>>": data["recomendaciones"] or "Sin recomendaciones adicionales.",
        "<<Juez>>": f"{data['juez_nombre']} (Cédula: {data['juez_cedula']})",
        "<<Marca temporal>>": str(data["fecha"])
    }

    reemplazar_etiquetas_word(doc, reemplazos)

    os.makedirs("descargas", exist_ok=True)
    output_filename = f"Evaluacion_{data['proyecto_codigo']}_{data['juez_cedula']}.docx"
    output_path = os.path.join("descargas", output_filename)
    doc.save(output_path)

    return FileResponse(
        path=output_path,
        filename=output_filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


@app.get("/proyecto/{proyecto_id}/acta-consolidada-pdf", response_class=HTMLResponse)
def vista_acta_consolidada(proyecto_id: int):
    conn = get_db()
    proyecto = conn.execute("SELECT * FROM proyectos WHERE id = ?", (proyecto_id,)).fetchone()

    if not proyecto:
        err = urllib.parse.quote("Proyecto no encontrado.")
        return RedirectResponse(url=f"/?error={err}", status_code=303)

    evaluaciones = conn.execute('''
        SELECT a.estado, j.nombre as juez_nombre, j.cedula as juez_cedula, e.nota_final, e.puntaje_total
        FROM asignaciones a
        JOIN jueces j ON a.juez_id = j.id
        LEFT JOIN evaluaciones e ON e.asignacion_id = a.id
        WHERE a.proyecto_id = ?
    ''', (proyecto_id,)).fetchall()
    conn.close()

    estudiantes = [proyecto['estudiante_1']]
    if proyecto['estudiante_2']: estudiantes.append(proyecto['estudiante_2'])
    if proyecto['estudiante_3']: estudiantes.append(proyecto['estudiante_3'])

    notas = [ev['nota_final'] for ev in evaluaciones if ev['nota_final'] is not None]
    promedio_general = round(sum(notas) / len(notas), 2) if notas else 0.0

    template = templates_env.get_template("plantilla_acta_consolidada_pdf.html")
    return template.render(
        proyecto=proyecto,
        estudiantes=estudiantes,
        evaluaciones=evaluaciones,
        promedio_general=promedio_general,
        nombre_institucion=obtener_institucion()
    )